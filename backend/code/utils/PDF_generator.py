from asyncio import sleep
from collections import Counter
import logging
import os
from datetime import datetime
from functools import reduce
import time
from utils.logger import logRequest
from docx import Document
from docx.shared import Pt
from babel.numbers import format_currency
from utils.calculadora import calcularArmazenagem, calcularLevante, calcularEnergia, calcularServicosConteiner
from utils.models import loadRecintos
import re

non_breaking_space = '\u00A0'

def calcularValoresContainer(tipo_mercadoria : str,data_entrada : datetime,data_saida : datetime,recinto : dict,valor_aduaneiro : float,containeres : int, servicos : list, tipo_container : str):
# Expand = True descreve o valor pago por período
    armazenagem = calcularArmazenagem(
        tipo_mercadoria=tipo_mercadoria,
        entrada = data_entrada,
        saida=data_saida,
        recinto=recinto,
        valor_aduaneiro=valor_aduaneiro,
        conteineres=containeres,
        expand=True
    )

    servicos = calcularServicosConteiner(servicos=servicos,recinto=recinto, tipo_mercadoria=tipo_mercadoria)

    levante = calcularLevante(tipo_container=tipo_container,recinto=recinto,tipo_mercadoria=tipo_mercadoria)

    energia =  calcularEnergia(entrada=data_entrada,saida=data_saida,recinto=recinto,tipo_mercadoria_input=tipo_mercadoria, expand=True)

    return armazenagem, servicos, levante, energia

# Defina value se o há apenas um valor, e listOfValues se for uma lista a ser concatenada com valor e descricao
def formatarValoresDescritos(template : str ,sequence : int, tipo : str, numero : str, entrada : str, saida : str, listOfValues : list):
    lista_containeres = map(lambda valor : template.format( 
        sequence=sequence,
        tipo=tipo,
        numero=numero,
        entrada=entrada,
        saida=saida,
        valor=str(round(valor["valor"],2)).ljust(12," "),
        descricao= valor["periodo"] + " periodo")+'\n', # Add a delimiter here for descritive purposes
        listOfValues)
    texto_formatado = ''.join(lista_containeres)
    return texto_formatado

def formatarValoresLevante(template : str, sequence : int, tipo : str, numero : str, entrada : str, saida : str, value : float):
    texto_formatado = template.format(sequence=sequence,
        tipo=tipo,
        numero=numero,
        entrada=entrada,
        saida=saida,
        valor=round(value,2),
        descricao= "")+'\n'
    return texto_formatado

def formatarValoresServicos(template : str, sequence : int, tipo : str, numero : str, entrada : str, saida : str, servicos : dict):
    # Agrupar todos os servicos em uma lista contendo o formato a baixo
    listServicos = {}
        # { "nome_servico" : nome, "texto" : texto_formatado}
    
    itemsServicos = servicos.items()
    for servico in itemsServicos:
        texto_formatado = template.format(sequence=sequence,
            tipo=tipo,
            numero=numero,
            entrada=entrada,
            saida=saida,
            valor=round(servico[1],2),
            descricao= "")+'\n'
        servico_formatado = { servico[0]: texto_formatado }
        listServicos = listServicos | servico_formatado
        
    
    return listServicos

def templateLoad() :
    document = Document("./utils/Template PDF.docx")
    return document

def generateKeyWords(data = None):

    recintos = loadRecintos()
    #Pega o recinto do modelo que foi passado
    recinto = recintos[data["recinto_nome"]]
    containeres = data["conteineres"]
    custo_obrigatorio = data['custo_obrigatorio']

    texto_armazenagem = ""
    lista_servicos = {}
    texto_levante = ""
    texto_energia = ""

    valor_armazenagem = 0
    valor_servicos = 0
    valor_levante = 0
    valor_energia = 0

    tab = "\u0009"
    texto_template = "{sequence} "+tab+"{tipo} "+tab+"{numero}"+tab+tab+"{entrada}"+tab+"{saida}"+tab+"{valor}"+tab+"{descricao}" 


    valor_aduaneiro = float(data["valor_aduaneiro"])
    for container in containeres:
        data_entrada = datetime.strptime(container["entrada"], '%d/%m/%Y')
        data_saida = datetime.strptime(container["saida"], '%d/%m/%Y')
        tipo_mercadoria = data["tipo_mercadoria"]

        servicos_input = custo_obrigatorio + container['servicos']

        armazenagem, servicos, levante, energia = calcularValoresContainer(tipo_mercadoria=tipo_mercadoria,
                                                                  data_entrada=data_entrada,
                                                                  data_saida=data_saida,
                                                                  recinto=recinto,
                                                                  valor_aduaneiro=valor_aduaneiro,
                                                                  containeres=containeres,
                                                                  servicos=servicos_input,
                                                                  tipo_container=container["tipo"])
       
        
        
        texto_armazenagem += formatarValoresDescritos(
            template=texto_template,
            sequence=container["sequence"],
            tipo=container["tipo"],
            numero=container["numero"],
            entrada=container["entrada"],
            saida=container["saida"],
            listOfValues=armazenagem
            ) 

        texto_levante += formatarValoresLevante(
            template=texto_template,
            sequence=container["sequence"],
            tipo=container["tipo"],
            numero=container["numero"],
            entrada=container["entrada"],
            saida=container["saida"],
            value=levante)

        servicos_formatados = formatarValoresServicos(
                template=texto_template,
                sequence=container["sequence"],
                tipo=container["tipo"],
                numero=container["numero"],
                entrada=container["entrada"],
                saida=container["saida"],
                servicos=servicos
            ) 
        
        
        if( len(energia) > 0 ):
            texto_energia += formatarValoresDescritos(
            template=texto_template,
            sequence=container["sequence"],
            tipo=container["tipo"],
            numero=container["numero"],
            entrada=container["entrada"],
            saida=container["saida"],
            listOfValues=energia
            ) 

        valor_armazenagem += reduce(lambda sum,valor : sum + valor["valor"],armazenagem,0)
        valor_energia += reduce(lambda sum,valor : sum + valor["valor"],energia,0)
        valor_servicos += reduce(lambda sum,valor : sum + valor[1],servicos.items(),0)
        valor_levante += levante
        # Concatena servicos formatados
        for servico in servicos_formatados.items():
            lista_servicos[servico[0]] = lista_servicos.get(servico[0],"") + servico[1]
    num_di = str(data['num_di'])
    formated_num_di = num_di[:2]+"/"+num_di[2:-1]+'-'+num_di[-1:]
    # Qualquer chave que for igual a uma tag no DOCX sera substituida. Caso seja necessário um comportamento especial, é necessário fazer uma chave específica
    return {
    "NUM_DI" :  formated_num_di,
    "REF_EXT" :  data['ref_ext'],
    "DATA_REGISTRO" :  data['data_registro'],
    "PTAX" :  str(data['ptax']),
    "VA" :  format_currency(data['valor_aduaneiro'],"BRL", locale="pt_BR"),
    "TIPO_MERCADORIA" :  data['tipo_mercadoria'],
    "RECINTO" :  data['recinto_nome'],
    "armazenagem" : texto_armazenagem,
    "levante" : texto_levante,
    "lista_servicos" : lista_servicos,
    "containers_energia" : texto_energia,
    "total_armazenagem" : format_currency(round(valor_armazenagem,2),"BRL", locale="pt_BR"),
    "total_levante" : format_currency(round(valor_levante,2),"BRL", locale="pt_BR"),
    "total_servicos" : format_currency(round(valor_servicos,2),"BRL", locale="pt_BR"), 
    "valor_energia" : format_currency(round(valor_energia,2),"BRL", locale="pt_BR"),
    "total_nota" : format_currency(round(valor_levante+valor_armazenagem+valor_servicos+valor_energia,2),"BRL", locale="pt_BR")
    }

def replace_match(match,keywords):
    # Non breaking space for justified texts
    # Extract the match (remove the curly braces)
    matched_key = str(match.group()[3:-3])  # Removes the '{{{' and '}}}' from the match
    return keywords.get(matched_key, match.group())  # Get the value or return the original match if not found

def replaceKeyWords(document, keywords = {}):
    key_list = keywords.keys()
    # Create the regular expression pattern dynamically with manual curly brace escape
    key_regex = re.compile(r"|".join(r"\{\{\{" + re.escape(substring) + r"\}\}\}" for substring in key_list))

    for paragraph in document.iter_inner_content() :
        #paragraph_font = paragraph.style.font.size
        texto_paragrafo = paragraph.text
        
        if(key_regex.search(texto_paragrafo)):
            logRequest(levelName=logging.INFO, message=texto_paragrafo)
            new_paragraph = key_regex.sub(lambda m: replace_match(m,keywords),texto_paragrafo)
            paragraph.text = new_paragraph
            #paragraph.font = paragraph_font
        
        if ("{{{energia}}}" in texto_paragrafo):
            if(len(keywords["containers_energia"]) > 0):
                logRequest(levelName=logging.INFO, message=texto_paragrafo)
                titulo_paragraph = paragraph.insert_paragraph_before()
                titulo_text = titulo_paragraph.add_run(
                    "Energia"
                )
                titulo_text.font.size = Pt(14)

                titulo_paragraph.paragraph_format.left_indent = -540385

                containers_paragraph = paragraph.insert_paragraph_before()
                containers_text = containers_paragraph.add_run(
                    keywords["containers_energia"]
                )
                containers_paragraph.paragraph_format.left_indent = -540385
                containers_paragraph.paragraph_format.right_indent = -940385
                containers_text.font.size = Pt(11)
            
            paragraph.text = None
        
        if ( "{{{total_energia}}}" in texto_paragrafo):
            if(len(keywords["containers_energia"]) > 0):
                new_paragraph = str(texto_paragrafo).replace("{{{total_energia}}}",keywords["valor_energia"])
                paragraph.text = new_paragraph
            else:
                paragraph.text = None


        # Get the servico containers template, add the services and then clean it
        if("{{{servicos}}}" in texto_paragrafo):
            logRequest(levelName=logging.INFO, message=texto_paragrafo)
            for servico in keywords["lista_servicos"].items():
                titulo_paragraph = paragraph.insert_paragraph_before()
                titulo_text = titulo_paragraph.add_run(
                    servico[0]
                )
                titulo_text.font.size = Pt(14)

                titulo_paragraph.paragraph_format.left_indent = -540385

                containers_paragraph = paragraph.insert_paragraph_before()
                containers_text = containers_paragraph.add_run(
                    servico[1]
                )
                containers_paragraph.paragraph_format.left_indent = -540385
                containers_text.font.size = Pt(11)
            
            paragraph.text = None




def generatePDF(data):
    document = templateLoad()
    keywords = generateKeyWords(data)

    replaceKeyWords(document=document, keywords=keywords)
    num_di = data["num_di"]
    document.save(f"./{num_di}.docx")
    time.sleep(0.5)
    os.system(f'libreoffice --headless --convert-to pdf {num_di}.docx')
    time.sleep(0.2)
    os.remove(f"{num_di}.docx")

